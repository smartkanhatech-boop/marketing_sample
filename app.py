import streamlit as st
import pandas as pd
from datetime import datetime, date
from fpdf import FPDF
import json
import os
import io
import base64
import math
import random

# Word Document Library
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional: Number to Words
try:
    from num2words import num2words
    HAS_NUM2WORDS = True
except ImportError:
    HAS_NUM2WORDS = False

# --- PAGE CONFIG ---
st.set_page_config(page_title="AD Billing Pro", layout="wide", page_icon="🏗️")

# --- AUTHENTICATION ---
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

def check_login():
    # Updated Credentials: ID: aman_giri_8962627817 password: smart_kanha_tech
    if st.session_state.username == 'aman_giri_8962627817' and st.session_state.password == 'smart_kanha_tech':
        st.session_state.authenticated = True
        st.session_state.login_error = False
    else:
        st.session_state.login_error = True

if not st.session_state.authenticated:
    st.markdown("<h2 style='text-align: center;'>🔒 AD Billing Login</h2>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password")
        st.button("Login", on_click=check_login, use_container_width=True)
        if st.session_state.get('login_error'):
            st.error("Access Denied")
    st.stop()

# --- CONSTANTS ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# Logo directly inside main folder instead of sub-folder
LOGO_FULL_PATH = os.path.join(BASE_DIR, "Logo.png")

# Updated Folder Name/DB path logic if needed, keeping structure same
DB_FILE = os.path.join(BASE_DIR, "ad_billing_db.json")

# Updated Company Details
COMPANY_NAME = "ABC SOLUTIONS PRIVATE LIMITED"
COMPANY_PHONE = "+91-0000000000 / 0000000000"
COMPANY_GSTIN = "GSTXXX151XXXXX"
COMPANY_ADDRESS = "123, BUSINESS PARK, METRO CITY (000000)"
BANK_DETAILS = {
    "bank": "SAMPLE PARTNER BANK",
    "ac": "0000123456789",
    "ifsc": "SAMP0000123",
    "name": "ABC SOLUTIONS PRIVATE LIMITED"
}

WORK_CATALOG = {
    "Site Visit": ["Initial site visit", "Client discussion", "Site supervision", "Chargeable Site Visit"],
    "Architecture & Design": ["Conceptual design", "Single line layout plan", "Double line layout plan", "Architectural working floor drawings", "Furniture layout", "Window Details", "Gate Details", "Elevation working", "Facade working details"],
    "Structural Design": ["Construction plan", "Column grid plan", "Footing detail", "Floor beams and slab details", "Column reinforcement", "Tank Details (Septic/Water)", "Compound wall details"],
    "Interior": ["Exterior drawing", "Working drawing"],
    "Electrical & Plumbing": ["Facade Electrical", "Conduit drawing", "Plumbing drawings", "Bathroom fitting details", "Chamber location", "Sewage details"],
    "2D & 3D": ["2D measurement drawing", "Facade design", "Facade 3D view"]
}

GST_RATES = {"0%": 0.00, "5%": 0.05, "12%": 0.12, "18%": 0.18}

# --- SESSION STATE INITIALIZATION ---
if 'invoice_data' not in st.session_state:
    st.session_state.invoice_data = {
        "items": [],
        "schedule": [],
        "meta": {"terms": """- 30% Advance prior to initiation of the work.
- 2 extra changes will be provided free of cost. Further changes will be chargeable as per requirement.
- Site visit will be chargeable (unless specified above).
- The project is to be completed within six months. In case of delay, the agreed price will be revised by 10% for every additional two months."""}
    }

if 'builder_c_name' not in st.session_state: st.session_state.builder_c_name = ""
if 'builder_c_mob' not in st.session_state: st.session_state.builder_c_mob = ""
if 'builder_c_addr' not in st.session_state: st.session_state.builder_c_addr = ""
if 'builder_dtype_idx' not in st.session_state: st.session_state.builder_dtype_idx = 0

if 'selected_cat' not in st.session_state:
    st.session_state.selected_cat = list(WORK_CATALOG.keys())[0]

if 'selected_descs' not in st.session_state:
    st.session_state.selected_descs = WORK_CATALOG[st.session_state.selected_cat]

if 'schedule_df' not in st.session_state:
    st.session_state.schedule_df = pd.DataFrame(columns=["Stage", "Amount", "Date"])

# --- CALLBACKS ---
def on_cat_change():
    new_cat = st.session_state.selected_cat
    if new_cat in WORK_CATALOG:
        st.session_state.selected_descs = WORK_CATALOG[new_cat]
    else:
        st.session_state.selected_descs = []

# --- DB & HELPERS ---
def save_db():
    with open(DB_FILE, 'w') as f: json.dump(st.session_state.db, f)

def load_db():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, 'r') as f: db = json.load(f)
        except: db = {"invoices": [], "quotations": [], "payments": []}
    else:
        db = {"invoices": [], "quotations": [], "payments": []}
    
    needs_save = False
    for rec in db.get('invoices', []):
        if 'id' not in rec:
            rec['id'] = f"INV-LEGACY-{random.randint(100,999)}"
            needs_save = True
        if 'status' not in rec:
            rec['status'] = "Pending"
            needs_save = True
            
    if 'payments' not in db:
        db['payments'] = []
        needs_save = True
        
    if needs_save:
        with open(DB_FILE, 'w') as f: json.dump(db, f)
    return db

if 'db' not in st.session_state:
    st.session_state.db = load_db()

def generate_next_id(doc_type):
    prefix = "INV" if doc_type == "FINAL BILL" else "QUOT"
    year = datetime.now().year
    records = st.session_state.db['invoices'] if prefix == "INV" else st.session_state.db['quotations']
    existing_seqs = []
    for r in records:
        parts = r.get('id', '').split('-')
        if len(parts) == 3 and parts[0] == prefix and parts[1] == str(year):
            try: existing_seqs.append(int(parts[2]))
            except: pass
    next_seq = max(existing_seqs) + 1 if existing_seqs else 1
    return f"{prefix}-{year}-{next_seq:03d}"

def sanitize_text(text):
    if not isinstance(text, str): text = str(text)
    text = text.replace('•', '-').replace('?', '-')
    text = text.replace('₹', 'Rs. ').replace('\u20b9', 'Rs. ')
    return text.encode('latin-1', 'replace').decode('latin-1')

def number_to_words_safe(amount):
    if HAS_NUM2WORDS:
        try:
            txt = num2words(amount, lang='en_IN').title()
            return sanitize_text(txt) + " Only"
        except: return "Check Amount"
    return f"{amount} (in words)"

def safe_float(val):
    try: return float(val)
    except: return 0.0

def calculate_totals(items, gst_rate_key):
    sub = 0.0
    for item in items:
        q = safe_float(item.get('qty', 0))
        r = safe_float(item.get('rate', 0))
        sub += q * r
    rate = GST_RATES.get(gst_rate_key, 0.0)
    gst = sub * rate
    grand = sub + gst
    return sub, gst, grand

# --- RECEIPT PDF ---
class ReceiptPDF(FPDF):
    def header(self): pass
    def footer(self): pass

def generate_receipt_bytes(payment_data):
    pdf = ReceiptPDF(format='A5', orientation='L')
    pdf.add_page()
    pdf.set_draw_color(0,0,0); pdf.rect(5, 5, 200, 138)
    if os.path.exists(LOGO_FULL_PATH): pdf.image(LOGO_FULL_PATH, 10, 10, 30)
    
    pdf.set_y(10); pdf.set_font('Times', 'B', 16); pdf.set_text_color(21, 101, 192)
    pdf.cell(0, 8, sanitize_text(COMPANY_NAME), 0, 1, 'R')
    pdf.set_text_color(0,0,0); pdf.set_font('Times', '', 9)
    pdf.cell(0, 5, sanitize_text(COMPANY_ADDRESS), 0, 1, 'R')
    pdf.cell(0, 5, sanitize_text(f"Ph: {COMPANY_PHONE}"), 0, 1, 'R')
    
    pdf.ln(10); pdf.set_font('Times', 'B', 14); pdf.cell(0, 10, "PAYMENT RECEIPT", 0, 1, 'C'); pdf.ln(5)
    
    pdf.set_font('Times', '', 12); pdf.set_x(20)
    pdf.write(8, "Received with thanks from  ")
    pdf.set_font('Times', 'B', 14); pdf.write(8, sanitize_text(payment_data['client_name']))
    pdf.set_font('Times', '', 12); pdf.write(8, "\n\n")
    
    pdf.set_x(20)
    amt = safe_float(payment_data['amount'])
    text = (f"The sum of  Rs. {amt:,.2f}\n"
            f"({number_to_words_safe(amt)})\n\n"
            f"Payment Date:  {payment_data['date']}\n"
            f"Payment Mode:  {payment_data['mode']}\n"
            f"Ref Invoice Date:  {payment_data.get('invoice_date', 'N/A')}")
    pdf.multi_cell(0, 8, sanitize_text(text))
    
    pdf.set_y(-30); pdf.set_font('Times', 'B', 10); pdf.cell(0, 5, "AUTHORIZED SIGNATORY", 0, 0, 'R')
    return pdf.output(dest='S').encode('latin-1')

# --- BILL PDF ---
class PDF(FPDF):
    def header(self): pass
    def footer(self): pass

def calculate_page_height(data, schedule_list):
    min_height = 297; required_height = 160 
    for item in data['items']:
        merged = item.get('merged_rows', [{'desc': item['desc']}])
        for r in merged:
            desc_len = len(r['desc'])
            lines = math.ceil(desc_len / 45) + r['desc'].count('\n')
            required_height += max(6, lines * 5)
    if schedule_list: required_height += 20 + (len(schedule_list) * 8)
    term_lines = data['meta']['terms'].count('\n') + 3
    required_height += (term_lines * 5)
    return max(min_height, required_height)

def generate_pdf_bytes(data, gst_rate_key, hide_gst, schedule_list):
    page_h = calculate_page_height(data, schedule_list)
    pdf = PDF(unit='mm', format=(210, page_h))
    pdf.set_auto_page_break(False); pdf.set_margins(10, 10, 10); pdf.add_page()
    
    if os.path.exists(LOGO_FULL_PATH): pdf.image(LOGO_FULL_PATH, 10, 10, 45)
    
    pdf.set_xy(110, 12); pdf.set_font('Times', 'B', 22); pdf.set_text_color(21, 101, 192)
    pdf.cell(90, 8, sanitize_text(COMPANY_NAME), 0, 1, 'R')
    pdf.set_xy(110, 20); pdf.set_text_color(0, 0, 0); pdf.set_font('Times', '', 10)
    pdf.cell(90, 5, sanitize_text(COMPANY_ADDRESS), 0, 1, 'R')
    pdf.set_xy(110, 25); pdf.cell(90, 5, sanitize_text(f"Ph: {COMPANY_PHONE}"), 0, 1, 'R')
    if not hide_gst: pdf.set_xy(110, 30); pdf.cell(90, 5, sanitize_text(f"GST: {COMPANY_GSTIN}"), 0, 1, 'R')
        
    pdf.set_draw_color(0, 0, 0); pdf.line(10, 45, 200, 45)
    
    pdf.set_y(52); pdf.set_font('Times', 'B', 16)
    display_type = "BILL" if data['meta']['type'] == "FINAL BILL" else data['meta']['type']
    pdf.cell(0, 8, sanitize_text(display_type), 0, 1, 'C')
    
    y_info = 65
    pdf.set_xy(10, y_info); pdf.set_font('Times', 'B', 10); pdf.cell(90, 5, "DOCUMENT DETAILS:", 0, 1)
    pdf.set_font('Times', '', 10); pdf.set_x(10)
    doc_id = data.get('id', 'NEW')
    pdf.cell(90, 5, sanitize_text(f"No: {doc_id}"), 0, 1)
    pdf.set_x(10); pdf.cell(90, 5, sanitize_text(f"Date: {data['meta']['date']}"), 0, 1)
    
    pdf.set_xy(110, y_info); pdf.set_font('Times', 'B', 10); pdf.cell(90, 5, "TO CLIENT:", 0, 1)
    pdf.set_xy(110, y_info + 6); pdf.set_font('Times', 'B', 12)
    pdf.cell(90, 6, sanitize_text(data['client']['name']), 0, 1)
    pdf.set_xy(110, y_info + 12); pdf.set_font('Times', '', 10)
    details = f"{data['client']['phone']}\n{data['client']['address']}"
    pdf.multi_cell(90, 5, sanitize_text(details))
    
    y_table_start = max(pdf.get_y(), y_info + 25) + 5
    pdf.set_xy(10, y_table_start)
    
    cols = [30, 80, 15, 15, 25, 25]
    headers = ["Category", "Description", "Unit", "Qty", "Rate (Rs.)", "Amount (Rs.)"]
    
    pdf.set_fill_color(240, 240, 240); pdf.set_font('Times', 'B', 10)
    for i, h in enumerate(headers):
        align = 'L' if i < 2 else 'R'
        pdf.cell(cols[i], 8, h, 1, 0, align, 1)
    pdf.ln()
    
    pdf.set_font('Times', '', 10)
    sub, gst, grand = calculate_totals(data['items'], gst_rate_key)
    if hide_gst: gst=0; grand=sub

    for item in data['items']:
        x_start = pdf.get_x()
        y_start = pdf.get_y()
        
        merged_rows = item.get('merged_rows', [])
        if not merged_rows:
            merged_rows = [{'cat': item['category'], 'desc': item['desc']}]
            
        total_h = 0
        sub_heights = []
        for sub_row in merged_rows:
            cat_txt = sanitize_text(sub_row['cat'])
            cat_lines = math.ceil(len(cat_txt) / 13) + cat_txt.count('\n')
            
            desc_txt = sanitize_text(sub_row['desc'])
            desc_lines = math.ceil(len(desc_txt) / 45) + desc_txt.count('\n')
            
            max_lines = max(cat_lines, desc_lines)
            h = max(6, max_lines * 5)
            
            sub_heights.append(h)
            total_h += h
            
        if total_h < 8: total_h = 8
        
        pdf.set_xy(x_start, y_start)
        pdf.rect(x_start, y_start, sum(cols), total_h)
        
        curr_x = x_start
        for w in cols[:-1]:
            curr_x += w
            pdf.line(curr_x, y_start, curr_x, y_start + total_h)
            
        curr_y = y_start
        for idx, sub_row in enumerate(merged_rows):
            h = sub_heights[idx]
            pdf.set_xy(x_start, curr_y)
            pdf.multi_cell(cols[0], 5, sanitize_text(sub_row['cat']), 0, 'L')
            pdf.set_xy(x_start + cols[0], curr_y)
            pdf.multi_cell(cols[1], 5, sanitize_text(sub_row['desc']), 0, 'L')
            if idx < len(merged_rows) - 1:
                line_y = curr_y + h
                pdf.line(x_start, line_y, x_start + cols[0] + cols[1], line_y)
            curr_y += h
            
        q = safe_float(item.get('qty', 0))
        r = safe_float(item.get('rate', 0))
        amt = q * r
        u = sanitize_text(item.get('unit', ''))
        
        pdf.set_xy(x_start + cols[0] + cols[1], y_start)
        pdf.cell(cols[2], total_h, u, 0, 0, 'R')
        pdf.set_xy(x_start + cols[0] + cols[1] + cols[2], y_start)
        pdf.cell(cols[3], total_h, f"{q}", 0, 0, 'R')
        pdf.set_xy(x_start + cols[0] + cols[1] + cols[2] + cols[3], y_start)
        pdf.cell(cols[4], total_h, f"{r:.2f}", 0, 0, 'R')
        pdf.set_xy(x_start + cols[0] + cols[1] + cols[2] + cols[3] + cols[4], y_start)
        pdf.cell(cols[5], total_h, f"{amt:.2f}", 0, 0, 'R')
        pdf.set_y(y_start + total_h)

    pdf.ln(2)
    def print_total(label, val, bold=False):
        if bold: pdf.set_font('Times', 'B', 11)
        else: pdf.set_font('Times', '', 10)
        v = safe_float(val)
        pdf.cell(140, 6, label, 0, 0, 'R'); pdf.cell(50, 6, f"Rs. {v:,.2f}", 0, 1, 'R')

    print_total("Subtotal:", sub)
    if not hide_gst: print_total(f"GST ({gst_rate_key}):", gst)
    print_total("Total:", grand, bold=True)
    pdf.ln(2); pdf.set_font('Times', 'I', 10)
    pdf.cell(0, 6, number_to_words_safe(grand), 0, 1, 'R')
    
    if schedule_list and len(schedule_list) > 0:
        pdf.ln(8); pdf.set_font('Times', 'B', 10); pdf.cell(0, 6, "PAYMENT SCHEDULE:", 0, 1, 'L')
        pdf.set_fill_color(245, 245, 245)
        pdf.cell(80, 6, "Stage", 1, 0, 'L', 1); pdf.cell(40, 6, "Amount", 1, 0, 'C', 1); pdf.cell(70, 6, "Date", 1, 1, 'L', 1)
        pdf.set_font('Times', '', 9)
        for r in schedule_list:
            pdf.cell(80, 6, sanitize_text(str(r.get("Stage",""))), 1)
            pdf.cell(40, 6, sanitize_text(str(r.get("Amount",""))), 1, 0, 'C')
            pdf.cell(70, 6, sanitize_text(str(r.get("Date",""))), 1, 1)

    pdf.ln(8); pdf.set_font('Times', 'B', 10); pdf.cell(0, 6, "TERMS & CONDITIONS:", 0, 1, 'L')
    pdf.set_font('Times', '', 10); pdf.multi_cell(0, 5, sanitize_text(data['meta']['terms']))
    
    pdf.ln(10); pdf.set_draw_color(0, 0, 0); pdf.line(10, pdf.get_y(), 200, pdf.get_y()); pdf.ln(5)
    y_foot = pdf.get_y(); pdf.set_font('Times', 'B', 10); pdf.cell(90, 5, "ACCOUNT DETAILS", 0, 1, 'L')
    pdf.set_font('Times', '', 9)
    pdf.cell(90, 5, sanitize_text(f"BANK: {BANK_DETAILS['bank']}"), 0, 1, 'L')
    pdf.cell(90, 5, sanitize_text(f"A/C: {BANK_DETAILS['ac']}"), 0, 1, 'L')
    pdf.cell(90, 5, sanitize_text(f"IFSC: {BANK_DETAILS['ifsc']}"), 0, 1, 'L')
    pdf.cell(90, 5, sanitize_text(f"NAME: {BANK_DETAILS['name']}"), 0, 1, 'L')
    pdf.set_xy(110, y_foot + 15); pdf.set_font('Times', 'B', 10); pdf.cell(88, 5, "AUTHORIZED SIGNATORY", 0, 0, 'R')
    return pdf.output(dest='S').encode('latin-1')

# --- DOCX GENERATOR ---
def generate_docx_bytes(data, gst_rate_key, hide_gst, schedule_list):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(10)
    ht = doc.add_table(rows=1, cols=2); ht.autofit = False; ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    if os.path.exists(LOGO_FULL_PATH): 
        try: ht.cell(0,0).paragraphs[0].add_run().add_picture(LOGO_FULL_PATH, width=Inches(2.0))
        except: pass
    p = ht.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(COMPANY_NAME + "\n"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(21, 101, 192)
    p.add_run(f"{COMPANY_ADDRESS}\nPh: {COMPANY_PHONE}")
    if not hide_gst: p.add_run(f"\nGST: {COMPANY_GSTIN}")
    doc.add_paragraph("_"*70)
    display_type = "BILL" if data['meta']['type'] == "FINAL BILL" else data['meta']['type']
    doc.add_paragraph(display_type).alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=1, cols=2); t.autofit = True
    doc_id = data.get('id', 'NEW')
    t.cell(0,0).paragraphs[0].add_run(f"DETAILS:\nType: {display_type}\nNo: {doc_id}\nDate: {data['meta']['date']}")
    c_cell = t.cell(0,1)
    c_cell.paragraphs[0].add_run("TO CLIENT:\n").bold = True
    c_cell.paragraphs[0].add_run(f"{data['client']['name']}\n").bold = True
    c_cell.paragraphs[0].add_run(f"{data['client']['phone']}\n{data['client']['address']}")
    doc.add_paragraph("\n")
    tbl = doc.add_table(rows=1, cols=6); tbl.style = 'Table Grid'
    hdrs = ["Category", "Description", "Unit", "Qty", "Rate (Rs.)", "Amount (Rs.)"]
    for i,h in enumerate(hdrs): tbl.rows[0].cells[i].text = h
    for item in data['items']:
        rc = tbl.add_row().cells
        rc[0].text=item['category']; rc[1].text=item['desc']
        q = safe_float(item.get('qty', 0)); r = safe_float(item.get('rate', 0)); u = str(item.get('unit', ''))
        rc[2].text=u; rc[3].text=f"{q}"; rc[4].text=f"{r:.2f}"; rc[5].text=f"{q*r:.2f}"
    sub, gst, grand = calculate_totals(data['items'], gst_rate_key)
    if hide_gst: gst=0; grand=sub
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"\nSubtotal: {sub:,.2f}\n"); 
    if not hide_gst: p.add_run(f"GST ({gst_rate_key}): {gst:,.2f}\n")
    p.add_run(f"Total: Rs. {grand:,.2f}").bold = True
    p.add_run(f"\n{number_to_words_safe(grand)}").italic = True
    if schedule_list:
        doc.add_paragraph("\nPAYMENT SCHEDULE:").runs[0].bold = True
        stbl = doc.add_table(rows=1, cols=3); stbl.style = 'Table Grid'; shdrs = ["Stage", "Amount", "Date"]
        for i, h in enumerate(shdrs): stbl.rows[0].cells[i].text = h
        for row in schedule_list:
            rc = stbl.add_row().cells; rc[0].text=str(row.get("Stage","")); rc[1].text=str(row.get("Amount","")); rc[2].text=str(row.get("Date",""))
    doc.add_paragraph("\nTERMS:\n"+data['meta']['terms']); doc.add_paragraph("_"*70)
    ft = doc.add_table(rows=1, cols=2); ft.autofit = True
    ft.cell(0,0).paragraphs[0].add_run(f"BANK DETAILS\nBANK: {BANK_DETAILS['bank']}\nA/C: {BANK_DETAILS['ac']}\nIFSC: {BANK_DETAILS['ifsc']}")
    ft.cell(0,1).paragraphs[0].add_run("\n\nAUTHORIZED SIGNATORY").bold = True
    f = io.BytesIO(); doc.save(f); f.seek(0); return f

# --- UI ---
st.title("🏗️ AD Billing Pro")

tab_b, tab_h, tab_t = st.tabs(["📝 Builder", "📂 History & Payments", "💰 Ledger"])

with tab_b:
    c1, c2, c3, c4 = st.columns(4)
    dtype = c1.selectbox("Type", ["QUOTATION", "FINAL BILL"], index=st.session_state.builder_dtype_idx)
    ddate = c2.date_input("Date", key="builder_date_picker")
    grate = c3.selectbox("GST", list(GST_RATES.keys()), index=3)
    hgst = c4.checkbox("Hide GST")
    
    st.markdown("---")
    cf, cp = st.columns([1, 1])
    
    with cf:
        c_name = st.text_input("Client Name", value=st.session_state.builder_c_name)
        c_mob = st.text_input("Client Mobile", value=st.session_state.builder_c_mob)
        c_addr = st.text_area("Client Address", value=st.session_state.builder_c_addr, height=60)
        
        st.session_state.builder_c_name = c_name
        st.session_state.builder_c_mob = c_mob
        st.session_state.builder_c_addr = c_addr
        
        st.subheader("Items")
        with st.container(border=True):
            item_mode = st.radio("Item Mode", ["Separate Rows (Single Category)", "Merged (Multiple Categories)"], horizontal=True)
            
            if item_mode == "Merged (Multiple Categories)":
                cats = st.multiselect("Select Categories", list(WORK_CATALOG.keys()))
                available_descs = []
                for c in cats: available_descs.extend(WORK_CATALOG[c])
                descs = st.multiselect("Descriptions", available_descs, default=available_descs)
            else:
                cat = st.selectbox("Category", list(WORK_CATALOG.keys()), key='selected_cat', on_change=on_cat_change)
                descs = st.multiselect("Description", WORK_CATALOG[cat], key='selected_descs')

            cust = st.text_input("Custom Desc.")
            c_q, c_r, c_u = st.columns(3)
            qty = c_q.number_input("Qty", 1.0); rate = c_r.number_input("Rate", 0.0, step=100.0); unit = c_u.selectbox("Unit", ["Sq.Ft", "Sq.Mt", "L/S", "Nos", "Job"])
            
            if st.button("➕ Add"):
                d_list = descs[:]
                if cust: d_list.append(cust)
                if d_list:
                    if item_mode == "Merged (Multiple Categories)":
                        merged_rows = []
                        cat_names = []
                        desc_names = []
                        for c in cats:
                            valid_descs = [d for d in d_list if d in WORK_CATALOG[c]]
                            if valid_descs:
                                merged_rows.append({'cat': c, 'desc': ", ".join(valid_descs)})
                                cat_names.append(c)
                                desc_names.append(f"{c} ({', '.join(valid_descs)})")
                        if cust:
                            merged_rows.append({'cat': 'Additional', 'desc': cust})
                            desc_names.append(f"Additional: {cust}")
                        
                        st.session_state.invoice_data['items'].append({
                            "category": "\n".join(cat_names), 
                            "desc": "\n".join(desc_names), 
                            "unit": unit, "qty": qty, "rate": rate,
                            "merged_rows": merged_rows
                        })
                    else:
                        st.session_state.invoice_data['items'].append({
                            "category": cat, "desc": ", ".join(d_list), "unit": unit, "qty": qty, "rate": rate,
                            "merged_rows": [{'cat': cat, 'desc': ", ".join(d_list)}]
                        })
                    st.rerun()
        
        if st.session_state.invoice_data['items']:
            st.dataframe(pd.DataFrame(st.session_state.invoice_data['items']))
            if st.button("Clear Items"): st.session_state.invoice_data['items'] = []; st.rerun()
            
        with st.expander("Payment Schedule (Optional)"):
            if 'Date' not in st.session_state.schedule_df.columns:
                st.session_state.schedule_df['Date'] = pd.Series(dtype='object')
            edited_sched = st.data_editor(st.session_state.schedule_df, num_rows="dynamic", use_container_width=True, column_config={"Date": st.column_config.DateColumn("Date", format="YYYY-MM-DD")})
            edited_sched['Date'] = edited_sched['Date'].apply(lambda x: x.strftime('%Y-%m-%d') if hasattr(x, 'strftime') else str(x) if pd.notnull(x) else "")
            st.session_state.schedule_df = edited_sched
            st.session_state.invoice_data['schedule'] = edited_sched.to_dict('records')

        term_txt = st.text_area("Terms", st.session_state.invoice_data['meta']['terms'], height=100)

    with cp:
        st.subheader("Live Preview")
        items = st.session_state.invoice_data['items']
        sub, gst, grand = calculate_totals(items, grate)
        if hgst: gst=0; grand=sub; gst_html=""
        else: gst_html=f"<tr><td colspan='5' align='right'>GST ({grate}):</td><td align='right'>Rs. {gst:,.2f}</td></tr>"
        
        logo_html = ""
        if os.path.exists(LOGO_FULL_PATH):
            with open(LOGO_FULL_PATH, "rb") as f: b64 = base64.b64encode(f.read()).decode()
            logo_html = f"<img src='data:image/png;base64,{b64}' width='120'>"
            
        rows_str = ""
        for i in items:
            rows_str += f"<tr><td>{i['category'].replace(chr(10),'<br>')}</td><td>{i['desc'].replace(chr(10),'<br>')}</td><td>{i.get('unit','')}</td><td align='right'>{i['qty']}</td><td align='right'>{i['rate']}</td><td align='right'>{i['qty']*i['rate']:.2f}</td></tr>"
        
        schedule_html = ""
        sched_data = [r for r in st.session_state.invoice_data.get('schedule',[]) if r.get("Stage") or r.get("Amount")]
        if sched_data:
            sch_rows = "".join([f"<tr><td>{r.get('Stage','')}</td><td>{r.get('Amount','')}</td><td>{r.get('Date','')}</td></tr>" for r in sched_data])
            schedule_html = f"""<div style="margin-top:15px; border:1px solid #ccc;"><strong>PAYMENT SCHEDULE:</strong><table style="width:100%; border-collapse:collapse; font-size:12px;"><tr style="background:#eee;"><th>Stage</th><th>Amount</th><th>Date</th></tr>{sch_rows}</table></div>"""

        display_type_html = "BILL" if dtype == "FINAL BILL" else dtype

        html = f"""<div style="border:1px solid #ddd; padding:20px; font-family:'Times New Roman'; color:black; background:white;">
<table style="width:100%; border:none;"><tr><td style="width:50%; vertical-align:top;">{logo_html}</td><td style="width:50%; text-align:right; vertical-align:top;"><h2 style="color:#1565C0; margin:0;">{COMPANY_NAME}</h2><div style="font-size:12px; color:black;">{COMPANY_ADDRESS}<br>Ph: {COMPANY_PHONE}</div></td></tr></table>
<hr style="border: 1px solid #333; margin: 10px 0;"><h3 style="text-align:center;">{display_type_html}</h3>
<table style="width:100%; border-collapse:collapse; margin-bottom:20px;"><tr><td style="width:48%; border:1px solid #ccc; padding:10px; vertical-align:top;"><strong>DETAILS:</strong><br>Type: {display_type_html}<br>Date: {ddate}</td><td style="width:4%; border:none;"></td><td style="width:48%; border:1px solid #ccc; padding:10px; vertical-align:top;"><strong>TO CLIENT:</strong><br><strong style="font-size:16px;">{c_name}</strong><br>{c_mob}<br>{c_addr}</td></tr></table>
<table style="width:100%; border-collapse:collapse; border:1px solid #ccc; font-size:13px;" border="1"><tr style="background:#eee;"><th>Cat</th><th>Desc</th><th>Unit</th><th>Qty</th><th>Rate</th><th>Amt</th></tr>{rows_str}<tr><td colspan='5' align='right'>Subtotal:</td><td align='right'>Rs. {sub:,.2f}</td></tr>{gst_html}<tr><td colspan='5' align='right'><b>Total:</b></td><td align='right'><b>Rs. {grand:,.2f}</b></td></tr></table>
<p style="text-align:right; font-style:italic;">{number_to_words_safe(grand)}</p>
{schedule_html}
<div style="border:1px dashed #ccc; padding:10px; margin-top:10px;"><strong>TERMS:</strong><pre style="white-space:pre-wrap; font-family:inherit; margin:0;">{term_txt}</pre></div>
<div style="margin-top:20px; border-top:2px solid black; padding-top:10px; display:flex; justify-content:space-between;"><div><strong>BANK DETAILS</strong><br>BANK: {BANK_DETAILS['bank']}<br>A/C: {BANK_DETAILS['ac']}<br>IFSC: {BANK_DETAILS['ifsc']}</div><div style="align-self:flex-end;"><strong>AUTHORIZED SIGNATORY</strong></div></div></div>"""
        
        st.markdown(html, unsafe_allow_html=True)
        
        next_id = generate_next_id(dtype)
        fdata = {"meta": {"type": dtype, "date": str(ddate), "terms": term_txt}, "client": {"name": c_name, "phone": c_mob, "address": c_addr}, "items": items, "id": next_id}
        
        st.divider()
        st.write("### 📄 Actions")
        c_dl1, c_dl2 = st.columns(2)
        
        with c_dl1:
            try:
                st.download_button("Download PDF", generate_pdf_bytes(fdata, grate, hgst, sched_data), f"{c_name}.pdf", "application/pdf")
            except Exception as e: st.error(f"PDF Error: {e}")
            
        with c_dl2:
            try:
                st.download_button("Download Word", generate_docx_bytes(fdata, grate, hgst, sched_data), f"{c_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e: st.error(f"Word Error: {e}")

        st.divider()
        if st.button("💾 Save & Finalize to History", type="primary", use_container_width=True):
            if c_name:
                rec = {
                    "id": next_id, "date": str(ddate), "type": dtype, 
                    "client_name": c_name, "client_phone": c_mob, "client_address": c_addr,
                    "amount": grand, "tax": gst, "items": items, "gst_rate": grate, "hide_gst": hgst,
                    "status": "Pending", "schedule": st.session_state.invoice_data['schedule'], "terms": term_txt
                }
                t = 'invoices' if dtype == "FINAL BILL" else 'quotations'
                st.session_state.db[t].append(rec); save_db(); st.toast(f"Saved: {next_id}")
            else: st.error("Name Required")

with tab_h:
    m = st.radio("View Mode", ["Quotations", "Bills & Payments"])
    
    if m == "Quotations":
        if st.session_state.db['quotations']:
            df_q = pd.DataFrame(st.session_state.db['quotations'])
            df_q['Summary'] = df_q['items'].apply(lambda x: ", ".join([i['category'].replace('\n', ', ') for i in x]))
            
            df_display = df_q.copy()
            df_display.index = df_display.index + 1
            
            st.dataframe(df_display[['id', 'date', 'client_name', 'amount', 'Summary']], use_container_width=True)
            
            sel_q_idx = st.selectbox("Select Quotation", range(len(df_q)), format_func=lambda x: f"{df_q.iloc[x]['client_name']} ({df_q.iloc[x]['amount']})")
            selected_quote = df_q.iloc[sel_q_idx]
            
            c1, c2, c3 = st.columns(3)
            if c1.button("✅ Confirm as Bill"):
                st.session_state.invoice_data['items'] = selected_quote['items']
                if 'schedule' in selected_quote:
                    st.session_state.invoice_data['schedule'] = selected_quote['schedule']
                    st.session_state.schedule_df = pd.DataFrame(selected_quote['schedule'])
                st.session_state.builder_c_name = selected_quote['client_name']
                st.session_state.builder_c_mob = selected_quote.get('client_phone', '')
                st.session_state.builder_c_addr = selected_quote.get('client_address', '')
                st.session_state.builder_dtype_idx = 1 
                del st.session_state.db['quotations'][sel_q_idx]
                save_db()
                st.success("Converted!"); st.rerun()
            
            if c2.button("✏️ Edit Quote"):
                st.session_state.invoice_data['items'] = selected_quote['items']
                st.session_state.builder_c_name = selected_quote['client_name']
                st.session_state.builder_c_mob = selected_quote.get('client_phone', '')
                st.session_state.builder_c_addr = selected_quote.get('client_address', '')
                del st.session_state.db['quotations'][sel_q_idx]
                save_db()
                st.success("Loaded!"); st.rerun()

            if c3.button("❌ Delete"):
                del st.session_state.db['quotations'][sel_q_idx]; save_db(); st.success("Deleted."); st.rerun()
                
            st.divider(); st.write("📄 Download Copy:")
            fdata_h = {
                "meta": {"type": selected_quote['type'], "date": selected_quote['date'], "terms": selected_quote.get('terms', st.session_state.invoice_data['meta']['terms'])},
                "client": {"name": selected_quote['client_name'], "phone": selected_quote.get('client_phone',''), "address": selected_quote.get('client_address','')},
                "items": selected_quote['items'], "id": selected_quote.get('id', 'N/A')
            }
            try: st.download_button("Download PDF", generate_pdf_bytes(fdata_h, selected_quote.get('gst_rate', '18%'), selected_quote.get('hide_gst', False), selected_quote.get('schedule', [])), f"{selected_quote['client_name']}.pdf", "application/pdf")
            except: st.error("Error generating PDF")
        else: st.info("No active quotations.")

    else: 
        if st.session_state.db['invoices']:
            df_i = pd.DataFrame(st.session_state.db['invoices'])
            def get_paid(iid): return sum([p['amount'] for p in st.session_state.db['payments'] if p.get('invoice_id') == iid])
            df_i['Paid'] = df_i['id'].apply(get_paid)
            df_i['Pending'] = df_i['amount'] - df_i['Paid']
            
            for idx, row in df_i.iterrows():
                if row['Pending'] <= 0 and row['status'] != "Completed":
                    st.session_state.db['invoices'][idx]['status'] = "Completed"; save_db()
            
            df_display = df_i.copy()
            df_display.index = df_display.index + 1
            
            st.dataframe(df_display[['id', 'date', 'client_name', 'amount', 'Paid', 'Pending', 'status']], use_container_width=True)
            st.divider()
            
            sel_b_idx = st.selectbox("Select Bill", range(len(df_i)), format_func=lambda x: f"{df_i.iloc[x]['client_name']} ({df_i.iloc[x]['status']})")
            db_record = st.session_state.db['invoices'][sel_b_idx]
            selected_bill_df = df_i.iloc[sel_b_idx]
            
            c1, c2 = st.columns(2)
            fdata_b = {
                "meta": {"type": db_record['type'], "date": db_record['date'], "terms": db_record.get('terms', "")},
                "client": {"name": db_record['client_name'], "phone": db_record.get('client_phone',''), "address": db_record.get('client_address','')},
                "items": db_record['items'], "id": db_record.get('id', 'N/A')
            }
            try: c1.download_button("📄 Download Bill PDF", generate_pdf_bytes(fdata_b, db_record.get('gst_rate','18%'), db_record.get('hide_gst', False), db_record.get('schedule',[])), f"{db_record['client_name']}.pdf", "application/pdf")
            except: c1.error("PDF Error")
            
            if db_record['status'] == "Pending":
                if c2.button("✅ Mark as Complete (Force)"):
                    st.session_state.db['invoices'][sel_b_idx]['status'] = "Completed (Manual)"; save_db(); st.rerun()

            if db_record['status'] == "Pending" and selected_bill_df['Pending'] > 1.0:
                st.subheader("💰 Record Payment")
                c1, c2, c3 = st.columns(3)
                pay_amt = c1.number_input("Amount", max_value=float(selected_bill_df['Pending']), value=float(selected_bill_df['Pending']), key="pay_amt")
                pay_date = c2.date_input("Date", key="pay_date_picker")
                pay_mode = c3.selectbox("Mode", ["UPI", "Cash", "Cheque", "Transfer"], key="pay_mode_sel")
                
                if st.button("Save Payment"):
                    p_rec = {
                        "id": f"PAY-{int(datetime.now().timestamp())}", "invoice_id": db_record['id'], "client_name": db_record['client_name'],
                        "invoice_date": db_record['date'], "amount": pay_amt, "date": str(pay_date), "mode": pay_mode
                    }
                    st.session_state.db['payments'].append(p_rec); save_db(); st.session_state.last_pay = p_rec
                    st.success("Payment Recorded!"); st.rerun()
                
                if 'last_pay' in st.session_state:
                    lp = st.session_state.last_pay
                    st.download_button("📄 Download Receipt", generate_receipt_bytes(lp), "Receipt.pdf", "application/pdf")
            else: st.info("Bill is Completed/Paid.")

with tab_t:
    st.header("Financial Ledger")
    c1, c2 = st.columns(2)
    start_d = c1.date_input("From Date", value=date(2024, 1, 1))
    end_d = c2.date_input("To Date", value=datetime.today())
    
    invs = st.session_state.db.get('invoices', [])
    pays = st.session_state.db.get('payments', [])
    
    f_invs = [i for i in invs if start_d <= datetime.strptime(i['date'], "%Y-%m-%d").date() <= end_d]
    f_pays = [p for p in pays if start_d <= datetime.strptime(p['date'], "%Y-%m-%d").date() <= end_d]
    
    t_billed = sum(i['amount'] for i in f_invs)
    t_rev = sum(p['amount'] for p in f_pays)
    t_gst = sum(i['tax'] for i in f_invs)
    
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total Billed", f"Rs. {t_billed:,.2f}")
    m2.metric("Revenue Received", f"Rs. {t_rev:,.2f}")
    m3.metric("Pending (Filtered)", f"Rs. {t_billed - t_rev:,.2f}")
    m4.metric("GST Collected", f"Rs. {t_gst:,.2f}")
    
    st.divider(); st.subheader("Client Reports")
    clients = sorted(list(set([i['client_name'] for i in invs])))
    sel_c = st.selectbox("Select Client", ["All"] + clients)
    
    if sel_c != "All":
        c_inv = [i for i in invs if i['client_name'] == sel_c]
        c_pay = [p for p in pays if p['client_name'] == sel_c]
        st.write(f"**Total Billed:** {sum(x['amount'] for x in c_inv):,.2f} | **Total Paid:** {sum(x['amount'] for x in c_pay):,.2f}")
        st.write("Bill History:"); st.dataframe(pd.DataFrame(c_inv))
        st.write("Payment History:"); st.dataframe(pd.DataFrame(c_pay))
        
    st.divider()
    c1, c2 = st.columns(2)
    if invs:
        df_exp = pd.DataFrame(invs)
        df_exp.index = df_exp.index + 1
        c1.download_button("📥 Export Bills CSV", df_exp.to_csv(), "bills.csv")
    if pays:
        df_rev = pd.DataFrame(pays)
        df_rev.index = df_rev.index + 1
        c2.download_button("📥 Export Revenue CSV", df_rev.to_csv(), "revenue.csv")